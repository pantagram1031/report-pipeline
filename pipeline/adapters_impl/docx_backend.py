# -*- coding: utf-8 -*-
"""docx_backend — the optional pure-python DOCX backend (`pip install .[docx]`).

Renders bundle/content.md to a styled .docx with python-docx. python-docx is an
OPTIONAL dependency: it is imported INSIDE build(), so importing this module
never fails on a plain clone. If it is absent, build() returns exit code 5 with
an install hint.

Mapping (bundle_spec.md build grammar → docx):
    title (build.yaml title:)  → document title paragraph
    ## SECTION:                 → Heading 1
    paragraph                   → body paragraph (inline EQ → italic code text)
    [[FIG file= width=NNmm]]    → embedded image (width mm, default 110mm)
    [[TABLE ...]] rows          → docx table (first row bold header)

Documented v1 limitation: equations are emitted as inline italic monospace text
(the literal latex/hwpeqn source), NOT as OMML — there is no equation typesetting
in the docx backend. PDF conversion is left to the user (LibreOffice:
`soffice --headless --convert-to pdf out.docx`).
"""
from __future__ import annotations

import os

from . import EQ_INLINE_RE, URL_INLINE_RE, eq_text, url_parts, parse_blocks, parse_front_matter

_DEFAULT_FIG_MM = 110


def _add_inline_runs(paragraph, text: str):
    """Append runs to a docx paragraph, rendering inline EQ/URL tags.

    EQ → italic monospace run carrying the literal source (no OMML in v1).
    URL → the display text as a plain run (no live hyperlink field in v1).
    """
    # Build an ordered token stream: (kind, payload) where kind in
    # {"text","eq","url"}. We walk EQ and URL matches together by position.
    spans = []
    for m in EQ_INLINE_RE.finditer(text):
        spans.append((m.start(), m.end(), "eq", eq_text(m.group(1))))
    for m in URL_INLINE_RE.finditer(text):
        spans.append((m.start(), m.end(), "url", url_parts(m.group(1))[1]))
    spans.sort()

    pos = 0
    for start, end, kind, payload in spans:
        if start < pos:
            continue  # overlap guard
        if start > pos:
            paragraph.add_run(text[pos:start])
        if kind == "eq":
            run = paragraph.add_run(payload)
            run.italic = True
            run.font.name = "Consolas"
        else:  # url
            paragraph.add_run(payload)
        pos = end
    if pos < len(text):
        paragraph.add_run(text[pos:])


def build(ws: str, out_dir: str | None = None):
    """Render bundle/content.md → <out_dir or WS/output>/out.docx.

    Returns (result: dict, exit_code: int):
      0  success
      2  bundle/content.md missing
      5  python-docx not installed (with install hint)
    """
    bundle = os.path.join(ws, "bundle")
    content_md = os.path.join(bundle, "content.md")
    if not os.path.isfile(content_md):
        return {"ok": False, "backend": "docx",
                "error": "bundle/content.md not found", "workspace": ws}, 2

    try:
        import docx
        from docx.shared import Mm, Pt
    except ImportError:
        return {"ok": False, "backend": "docx",
                "error": "python-docx not installed",
                "hint": "pip install python-docx"}, 5

    md = open(content_md, encoding="utf-8").read()
    meta, body = parse_front_matter(md)
    blocks = parse_blocks(body)

    build_yaml = os.path.join(ws, "build.yaml")
    title = None
    if os.path.isfile(build_yaml):
        from . import read_build_yaml_key
        title = read_build_yaml_key(build_yaml, "title")
    if not title:
        title = meta.get("title")

    doc = docx.Document()
    if title:
        doc.add_heading(title, level=0)

    figdir = os.path.join(bundle, "figures")
    for b in blocks:
        t = b["type"]
        if t == "section":
            doc.add_heading(b["title"], level=1)
        elif t == "para":
            if b["text"]:
                _add_inline_runs(doc.add_paragraph(), b["text"])
        elif t == "fig":
            img = os.path.join(figdir, b["file"].replace("\\", "/").lstrip("/"))
            # Defense in depth (content_audit already rejects traversal): never
            # embed an image resolving outside bundle/figures.
            fig_real = os.path.realpath(figdir)
            try:
                img_ok = os.path.commonpath([fig_real, os.path.realpath(img)]) == fig_real
            except ValueError:
                img_ok = False
            if not img_ok:
                p = doc.add_paragraph()
                p.add_run(f"[figure outside bundle/figures rejected: {b['file']}]").italic = True
                continue
            width_mm = b["width"] or _DEFAULT_FIG_MM
            p = doc.add_paragraph()
            p.alignment = 1  # center
            if os.path.isfile(img):
                run = p.add_run()
                run.add_picture(img, width=Mm(width_mm))
            else:
                p.add_run(f"[missing figure: {b['file']}]").italic = True
            if b["caption"]:
                cap = doc.add_paragraph(b["caption"])
                cap.alignment = 1
                for r in cap.runs:
                    r.font.size = Pt(9)
        elif t == "table":
            rows = b["rows"]
            if not rows:
                continue
            ncols = max(len(r) for r in rows)
            table = doc.add_table(rows=0, cols=ncols)
            table.style = "Table Grid"
            for ri, row in enumerate(rows):
                cells = table.add_row().cells
                for ci in range(ncols):
                    val = row[ci] if ci < len(row) else ""
                    cells[ci].text = val
                    if ri == 0:
                        for para in cells[ci].paragraphs:
                            for run in para.runs:
                                run.bold = True

    dest = out_dir or os.path.join(ws, "output")
    os.makedirs(dest, exist_ok=True)
    out_path = os.path.join(dest, "out.docx")
    doc.save(out_path)

    result = {
        "ok": True,
        "backend": "docx",
        "workspace": ws,
        "out": out_path,
        "title": title,
        "figures": sum(1 for b in blocks if b["type"] == "fig"),
        "tables": sum(1 for b in blocks if b["type"] == "table"),
        "sections": sum(1 for b in blocks if b["type"] == "section"),
        "limitation": "equations rendered as inline text (no OMML in v1)",
    }
    return result, 0
